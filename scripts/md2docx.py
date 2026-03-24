#!/usr/bin/env python3
"""
Markdown to Word Document Converter
====================================

将Markdown文件转换为Word文档，支持中文、表格、富文本格式。

特性:
    - 完整格式解析：加粗、斜体、代码、链接、删除线
    - 富文本支持：段落和表格中正确应用格式
    - 单元格换行：<br>自动转换

用法:
    python3 md2docx.py input.md output.docx [options]

作者: CoPaw AI Assistant
版本: 1.3.0
"""

import sys
import re
import argparse
import html
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class RichTextParser:
    """富文本解析器 - 解析Markdown并生成带格式的runs"""
    
    @staticmethod
    def parse_to_runs(text, font_name='微软雅黑', font_size=11, base_bold=False, base_italic=False):
        """
        解析文本并生成带格式的runs列表
        
        返回: [(文本, 加粗, 斜体, 代码, 删除线), ...]
        """
        if not text:
            return []
        
        # 处理HTML实体
        text = html.unescape(text)
        # 处理<br>换行
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        # 移除其他HTML标签
        text = re.sub(r'</?[a-zA-Z][^>]*>', '', text)
        
        # 解析格式（传入基础格式）
        segments = RichTextParser._parse_formatting(text, base_bold, base_italic)
        
        return segments
    
    @staticmethod
    def _parse_formatting(text, base_bold=False, base_italic=False):
        """
        解析文本中的格式标记，返回带格式信息的片段列表
        
        使用栈式解析器处理嵌套格式
        """
        result = []
        i = 0
        n = len(text)
        
        while i < n:
            # 检查链接 [...](...)
            if text[i] == '[':
                end_bracket = text.find('](', i)
                if end_bracket != -1:
                    end_paren = text.find(')', end_bracket)
                    if end_paren != -1:
                        # 提取链接文字
                        link_text = text[i+1:end_bracket]
                        # 链接文字可能有内部格式，递归解析
                        inner = RichTextParser._parse_formatting(link_text, base_bold, base_italic)
                        result.extend(inner)
                        i = end_paren + 1
                        continue
            
            # 检查删除线 ~~...~~
            if text[i:i+2] == '~~':
                end = text.find('~~', i+2)
                if end != -1:
                    inner_text = text[i+2:end]
                    inner = RichTextParser._parse_formatting(inner_text, base_bold, base_italic)
                    for seg_text, bold, italic, code, _ in inner:
                        result.append((seg_text, bold, italic, code, True))
                    i = end + 2
                    continue
            
            # 检查加粗 **...** 或 __...__
            if text[i:i+2] in ('**', '__'):
                marker = text[i:i+2]
                end = text.find(marker, i+2)
                if end != -1:
                    inner_text = text[i+2:end]
                    # 递归解析，设置基础加粗
                    inner = RichTextParser._parse_formatting(inner_text, True, base_italic)
                    result.extend(inner)
                    i = end + 2
                    continue
            
            # 检查斜体 *...* 或 _..._（单个符号）
            if text[i] in ('*', '_'):
                marker = text[i]
                # 确保不是加粗标记的一部分
                if i+1 < n and text[i+1] == marker:
                    # 这是加粗的开始，跳过让上面的加粗处理
                    pass
                else:
                    # 查找结束的斜体标记
                    end = i + 1
                    while end < n:
                        if text[end] == marker and (end+1 >= n or text[end+1] != marker):
                            break
                        end += 1
                    
                    if end < n:
                        inner_text = text[i+1:end]
                        inner = RichTextParser._parse_formatting(inner_text, base_bold, True)
                        result.extend(inner)
                        i = end + 1
                        continue
            
            # 检查代码 `...`
            if text[i] == '`':
                end = text.find('`', i+1)
                if end != -1:
                    code_text = text[i+1:end]
                    result.append((code_text, base_bold, base_italic, True, False))
                    i = end + 1
                    continue
            
            # 普通字符，累积到当前片段
            result.append((text[i], base_bold, base_italic, False, False))
            i += 1
        
        # 合并相邻的相同格式片段
        merged = []
        for seg in result:
            if merged and merged[-1][1:] == seg[1:]:
                # 格式相同，合并文本
                merged[-1] = (merged[-1][0] + seg[0],) + seg[1:]
            else:
                merged.append(seg)
        
        return merged
    
    @staticmethod
    def clean_text(text):
        """清理所有Markdown标记，返回纯文本"""
        if not text:
            return ''
        
        result = text
        result = html.unescape(result)
        result = re.sub(r'<br\s*/?>', '\n', result, flags=re.IGNORECASE)
        result = re.sub(r'</?[a-zA-Z][^>]*>', '', result)
        result = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', result)
        result = re.sub(r'\*\*([^*]+)\*\*', r'\1', result)
        result = re.sub(r'__([^_]+)__', r'\1', result)
        result = re.sub(r'(?<!\*)\*(?!\*)([^*]+)(?<!\*)\*(?!\*)', r'\1', result)
        result = re.sub(r'(?<!_)_(?!_)([^_]+)(?<!_)_(?!_)', r'\1', result)
        result = re.sub(r'`([^`]+)`', r'\1', result)
        result = re.sub(r'~~([^~]+)~~', r'\1', result)
        result = re.sub(r' +', ' ', result)
        
        return result.strip()


class MarkdownToDocx:
    """Markdown转Word转换器"""
    
    def __init__(self, font_name='微软雅黑', font_size=11, title_size=28, page_break=True):
        self.font_name = font_name
        self.font_size = font_size
        self.title_size = title_size
        self.page_break = page_break
        self.doc = Document()
        self._set_default_font()
    
    def _set_default_font(self):
        """设置文档默认字体"""
        style = self.doc.styles['Normal']
        style.font.name = self.font_name
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        style.font.size = Pt(self.font_size)
    
    def _set_run_font(self, run, bold=False, italic=False, font_size=None, color=None):
        """设置run的字体样式"""
        run.font.size = Pt(font_size or self.font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        run.font.name = self.font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
    
    def _set_cell_shading(self, cell, color):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _add_rich_text(self, paragraph, text, base_bold=False, base_italic=False):
        """添加富文本到段落"""
        runs = RichTextParser.parse_to_runs(text, base_bold=base_bold, base_italic=base_italic)
        
        for run_text, bold, italic, code, strike in runs:
            run = paragraph.add_run(run_text)
            self._set_run_font(run, bold=bold, italic=italic)
            if code:
                run.font.name = 'Courier New'
            if strike:
                run.font.strike = True
    
    def convert(self, md_content):
        """转换Markdown内容为Word文档"""
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if not line:
                i += 1
                continue
            
            # 标题
            if line.startswith('#'):
                self._add_heading(line)
                i += 1
                continue
            
            # 表格
            if line.startswith('|') and '|' in line[1:]:
                table_data, consumed = self._parse_table(lines, i)
                self._add_table(table_data)
                i += consumed
                continue
            
            # 代码块
            if line.startswith('```'):
                code_block, consumed = self._parse_code_block(lines, i)
                self._add_code_block(code_block)
                i += consumed
                continue
            
            # 引用块
            if line.startswith('>'):
                text = line[1:].strip()
                self._add_quote(text)
                i += 1
                continue
            
            # 无序列表
            if line.startswith(('- ', '* ', '• ')):
                items, consumed = self._parse_list(lines, i, ordered=False)
                self._add_list(items, ordered=False)
                i += consumed
                continue
            
            # 有序列表
            if re.match(r'^\d+\.\s', line):
                items, consumed = self._parse_list(lines, i, ordered=True)
                self._add_list(items, ordered=True)
                i += consumed
                continue
            
            # 分隔线
            if line in ('---', '***', '___'):
                self.doc.add_paragraph('─' * 50)
                i += 1
                continue
            
            # 普通段落（支持富文本）
            para = self.doc.add_paragraph()
            self._add_rich_text(para, line)
            i += 1
        
        return self.doc
    
    def _add_heading(self, line):
        """添加标题"""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if not match:
            return
        
        level = len(match.group(1))
        text = match.group(2)
        
        if self.page_break and level == 1 and len(self.doc.paragraphs) > 0:
            self.doc.add_page_break()
        
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            font_sizes = {1: 24, 2: 18, 3: 14, 4: 12, 5: 11, 6: 10}
            self._set_run_font(run, bold=True, font_size=font_sizes.get(level, self.font_size))
            # 设置标题颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _parse_table(self, lines, start):
        """解析表格"""
        table_data = []
        i = start
        
        while i < len(lines) and lines[i].startswith('|'):
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            
            if all(set(cell) <= set('-:|') for cell in row):
                i += 1
                continue
            
            table_data.append(row)
            i += 1
        
        return table_data, i - start
    
    def _add_table(self, table_data):
        """添加表格（支持富文本）"""
        if not table_data:
            return
        
        num_cols = max(len(row) for row in table_data)
        table = self.doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    
                    # 使用富文本解析
                    is_header = (row_idx == 0)
                    self._add_rich_text(p, cell_data, base_bold=is_header)
                    
                    if is_header:
                        self._set_cell_shading(cell, 'E7E6E6')
    
    def _parse_list(self, lines, start, ordered=False):
        """解析列表"""
        items = []
        i = start
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if ordered:
                if re.match(r'^\d+\.\s', line):
                    items.append(re.sub(r'^\d+\.\s+', '', line))
                    i += 1
                else:
                    break
            else:
                if line.startswith(('- ', '* ', '• ')):
                    items.append(line[2:].strip())
                    i += 1
                else:
                    break
        
        return items, i - start
    
    def _add_list(self, items, ordered=False):
        """添加列表（支持富文本）"""
        for idx, item in enumerate(items, 1):
            prefix = f'{idx}. ' if ordered else '• '
            para = self.doc.add_paragraph()
            run = para.add_run(prefix)
            self._set_run_font(run)
            self._add_rich_text(para, item)
    
    def _parse_code_block(self, lines, start):
        """解析代码块"""
        code_lines = []
        i = start + 1
        
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        return '\n'.join(code_lines), i - start + 1
    
    def _add_code_block(self, code):
        """添加代码块"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    def _add_quote(self, text):
        """添加引用块（支持富文本）"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        
        # 判断是否为来源
        if text.startswith('来源') or text.startswith('Source'):
            run = para.add_run(text)
            run.italic = True
            self._set_run_font(run, font_size=10)
        else:
            run = para.add_run('"')
            self._set_run_font(run)
            self._add_rich_text(para, text)
            run = para.add_run('"')
            self._set_run_font(run)
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        return output_path


def main():
    parser = argparse.ArgumentParser(description='Markdown转Word（支持富文本格式）')
    parser.add_argument('input', help='输入Markdown文件')
    parser.add_argument('output', help='输出Word文件')
    parser.add_argument('--font', default='微软雅黑', help='默认字体')
    parser.add_argument('--size', type=int, default=11, help='默认字号')
    parser.add_argument('--title-size', type=int, default=28, help='标题字号')
    parser.add_argument('--no-page-break', action='store_true', help='禁用章节分页')
    
    args = parser.parse_args()
    
    md_path = Path(args.input)
    if not md_path.exists():
        print(f'错误: 文件不存在 - {args.input}')
        sys.exit(1)
    
    md_content = md_path.read_text(encoding='utf-8')
    
    converter = MarkdownToDocx(
        font_name=args.font,
        font_size=args.size,
        title_size=args.title_size,
        page_break=not args.no_page_break
    )
    converter.convert(md_content)
    converter.save(args.output)
    
    print(f'转换完成: {args.output}')


if __name__ == '__main__':
    main()
