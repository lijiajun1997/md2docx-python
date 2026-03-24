#!/usr/bin/env python3
"""
Structured Report Template for Word Document Generation
========================================================

结构化报告生成模板，适用于已有结构化数据的报告生成场景。

使用方式:
1. 复制此脚本到工作目录
2. 修改 DATA 部分的报告数据
3. 运行脚本生成报告

作者: CoPaw AI Assistant
版本: 1.0.0
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# 样式工具函数
# ============================================================================

def set_run_font(run, font_name='微软雅黑', font_size=11, bold=False, color=None):
    """
    设置run的字体样式
    
    参数:
        run: docx run对象
        font_name: 字体名称（默认：微软雅黑）
        font_size: 字号（默认：11pt）
        bold: 是否加粗
        color: 颜色（RGBColor对象）
    """
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 关键：同时设置中文字体
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_shading(cell, color):
    """
    设置单元格背景色
    
    参数:
        cell: 表格单元格对象
        color: 16进制颜色代码（如 'E7E6E6'）
    """
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_color='E7E6E6'):
    """
    添加带样式的表格
    
    参数:
        doc: 文档对象
        headers: 表头列表
        rows: 数据行列表（每个元素是一行数据）
        header_color: 表头背景色
    
    返回:
        table对象
    """
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(header))
        set_run_font(run, bold=True)
        set_cell_shading(cell, header_color)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx+1, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_data))
            set_run_font(run)
    
    return table


def add_heading(doc, text, level=1):
    """
    添加标题
    
    参数:
        doc: 文档对象
        text: 标题文本
        level: 标题级别（1-6）
    
    返回:
        heading对象
    """
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        font_sizes = {1: 24, 2: 18, 3: 14, 4: 12, 5: 11, 6: 10}
        set_run_font(run, bold=True, font_size=font_sizes.get(level, 11))
    return h


def add_para(doc, text, bold_prefix=''):
    """
    添加段落
    
    参数:
        doc: 文档对象
        text: 段落文本
        bold_prefix: 加粗前缀（可选）
    
    返回:
        paragraph对象
    """
    p = doc.add_paragraph()
    if bold_prefix:
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, bold=True)
        run2 = p.add_run(text)
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_cover(doc, title, subtitle='', info_lines=None):
    """
    添加封面
    
    参数:
        doc: 文档对象
        title: 主标题
        subtitle: 副标题
        info_lines: 底部信息行列表
    """
    doc.add_paragraph()
    
    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    set_run_font(run, bold=True, font_size=28, color=RGBColor(0x2F, 0x54, 0x96))
    
    # 副标题
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        set_run_font(run, bold=True, font_size=28, color=RGBColor(0x2F, 0x54, 0x96))
    
    doc.add_paragraph()
    
    # 底部信息
    if info_lines:
        for line in info_lines:
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run(line)
            set_run_font(run, font_size=12)


def add_source(doc, text):
    """添加来源引用（斜体小字）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_size=10)
    run.italic = True
    return p


# ============================================================================
# 报告数据（请修改此处）
# ============================================================================

REPORT_DATA = {
    'title': '报告标题',
    'subtitle': '副标题（可选）',
    'info_lines': [
        '作者：XXX',
        '日期：2026年3月24日',
        '版本：v1.0'
    ],
    'sections': [
        {
            'heading': '第一章 概述',
            'level': 1,
            'content': [
                {
                    'type': 'para',
                    'text': '这是一个示例段落。段落可以包含普通文本。'
                },
                {
                    'type': 'para',
                    'text': '这个段落有加粗前缀。',
                    'bold_prefix': '要点：'
                },
                {
                    'type': 'table',
                    'headers': ['列1', '列2', '列3'],
                    'rows': [
                        ['数据1', '数据2', '数据3'],
                        ['数据4', '数据5', '数据6'],
                    ]
                },
                {
                    'type': 'heading',
                    'text': '1.1 子章节',
                    'level': 2
                },
                {
                    'type': 'source',
                    'text': '来源：示例数据源'
                }
            ]
        }
    ]
}


# ============================================================================
# 报告生成器
# ============================================================================

def generate_report(data, output_path):
    """
    根据数据生成报告
    
    参数:
        data: 报告数据字典
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    
    # 封面
    add_cover(doc, data['title'], data.get('subtitle', ''), data.get('info_lines'))
    doc.add_page_break()
    
    # 章节
    for section in data.get('sections', []):
        add_heading(doc, section['heading'], section.get('level', 1))
        
        for item in section.get('content', []):
            item_type = item.get('type', 'para')
            
            if item_type == 'para':
                add_para(doc, item['text'], item.get('bold_prefix', ''))
            
            elif item_type == 'heading':
                add_heading(doc, item['text'], item.get('level', 2))
            
            elif item_type == 'table':
                add_table(doc, item['headers'], item['rows'])
            
            elif item_type == 'source':
                add_source(doc, item['text'])
    
    # 保存
    doc.save(output_path)
    return output_path


# ============================================================================
# 主程序
# ============================================================================

if __name__ == '__main__':
    output_file = 'generated_report.docx'
    generate_report(REPORT_DATA, output_file)
    print(f'报告已生成: {output_file}')
